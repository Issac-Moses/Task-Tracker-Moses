import os
import json
import whisper
import sounddevice as sd
from scipy.io.wavfile import write
from transformers import pipeline
from fpdf import FPDF
from docx import Document
import yt_dlp as youtube_dl

# -------------- Global Configs --------------
samplerate = 44100
summarizer_pipeline = pipeline("summarization", model="google/flan-t5-large")

# -------------- Audio Recording (Mic) --------------
def record_audio(duration=5, filename="mic_output.wav"):
    print("🎙️ Recording Started...")
    try:
        audio = sd.rec(int(duration * samplerate), samplerate=samplerate, channels=1, dtype='float32')
        sd.wait()
        write(filename, samplerate, (audio * 32767).astype('int16'))
        print("✅ Recording Saved as", filename)
        return filename
    except Exception as e:
        print("❌ Error while recording:", e)
        return None

# -------------- Transcription Using Whisper --------------
def transcribe_audio(file_path):
    try:
        model = whisper.load_model("medium")
        result = model.transcribe(file_path)
        return result["text"]
    except Exception as e:
        raise RuntimeError(f"Transcription failed: {str(e)}")

def transcribe_youtube(youtube_url):
    try:
        print("⬇️ Downloading YouTube audio...")
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': 'yt_audio.%(ext)s',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }]
        }
        with youtube_dl.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(youtube_url, download=True)
            filename = ydl.prepare_filename(info).replace(".webm", ".mp3").replace(".m4a", ".mp3")
        print("✅ Audio downloaded:", filename)
        return transcribe_audio(filename), filename
    except Exception as e:
        raise RuntimeError(f"YouTube transcription failed: {str(e)}")

# -------------- Chunking Text --------------
def chunk_text(text, max_words=1000):
    sentences = text.split('. ')
    chunks = []
    current_chunk = ""
    for sentence in sentences:
        if len(current_chunk.split()) + len(sentence.split()) < max_words:
            current_chunk += sentence + ". "
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence + ". "
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

# -------------- Summarization --------------
def generate_summary(text):
    chunks = chunk_text(text)
    summaries = [summarizer_pipeline(chunk, max_length=350, min_length=30, do_sample=False)[0]['summary_text'] for chunk in chunks]
    return "\n".join(summaries)

def generate_key_points(text):
    prompt = "Extract key points as bullet points:\n\n" + text
    chunks = chunk_text(prompt)
    keypoints = [summarizer_pipeline(chunk, max_length=400, min_length=50, do_sample=False)[0]['summary_text'] for chunk in chunks]
    return "\n".join(keypoints)

# -------------- Export Options --------------
def export_to_pdf(summary, overview, keypoints):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Lecture Summary", ln=True, align='C')
    pdf.multi_cell(0, 10, "Overall Summary:\n" + summary + "\n\n")
    pdf.multi_cell(0, 10, "Overview:\n" + overview + "\n\n")
    pdf.multi_cell(0, 10, "Key Points:\n" + keypoints + "\n\n")
    pdf.output("lecture_summary.pdf")

def export_to_word(summary, overview, keypoints):
    doc = Document()
    doc.add_heading("Lecture Summary", 0)
    doc.add_heading("Overall Summary:", level=1)
    doc.add_paragraph(summary)
    doc.add_heading("Overview:", level=1)
    doc.add_paragraph(overview)
    doc.add_heading("Key Points:", level=1)
    doc.add_paragraph(keypoints)
    doc.save("lecture_summary.docx")

def export_to_json(summary, overview, keypoints):
    with open("lecture_summary.json", "w") as f:
        json.dump({
            "overall_summary": summary,
            "overview": overview,
            "key_points": keypoints
        }, f)

# -------------- Main Processing Function --------------
def process_input(source_type="mic", file_path=None, youtube_url=None, duration=10, export_format="PDF"):
    try:
        cleanup_files = []
        if source_type == "mic":
            file_path = record_audio(duration=duration)
            transcript = transcribe_audio(file_path)
            cleanup_files.append(file_path)
        elif source_type == "file" and file_path:
            transcript = transcribe_audio(file_path)
        elif source_type == "youtube" and youtube_url:
            transcript, temp_file = transcribe_youtube(youtube_url)
            cleanup_files.append(temp_file)
        else:
            return {"error": "Invalid input source."}

        summary = generate_summary(transcript)
        overview = generate_summary(transcript[:1000])
        keypoints = generate_key_points(transcript)

        output_file = ""
        export_format = export_format.upper()
        if export_format == "PDF":
            export_to_pdf(summary, overview, keypoints)
            output_file = "lecture_summary.pdf"
        elif export_format == "WORD":
            export_to_word(summary, overview, keypoints)
            output_file = "lecture_summary.docx"
        elif export_format == "JSON":
            export_to_json(summary, overview, keypoints)
            output_file = "lecture_summary.json"
        else:
            return {"error": "Unsupported export format."}

        for f in cleanup_files:
            if os.path.exists(f): os.remove(f)

        return {
            "overall_summary": summary,
            "overview": overview,
            "keypoints": keypoints,
            "output_file": output_file
        }

    except Exception as e:
        return {"error": str(e)}
